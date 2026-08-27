from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
POST = HERE / "SECTION_25_2_DISCUSSION_BOARD_POST.md"
OUT = HERE / "BBO_25_2_Successful_Optimisation_Strategies.docx"
PAIRS = HERE / "figure_pairs"
FIGURES = HERE / "figures"
BBR_PAIRS = HERE / "bbr_pairs"

NAVY = "102A43"
BLUE = "2F6B9A"
TEAL = "2A9D8F"
MUTED = "64748B"


def set_font(run, name="Aptos", size=10.2, colour=None, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if colour:
        run.font.color.rgb = RGBColor.from_string(colour)


def add_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.widow_control = True
    parts = re.split(r"(`[^`]+`|https://\S+)", text)
    for piece in parts:
        if not piece:
            continue
        if piece.startswith("`"):
            r = p.add_run(piece[1:-1]); set_font(r, "Consolas", 9.5, NAVY)
        elif piece.startswith("https://"):
            hyperlink = OxmlElement("w:hyperlink")
            rel = p.part.relate_to(piece, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            hyperlink.set(qn("r:id"), rel)
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            colour = OxmlElement("w:color"); colour.set(qn("w:val"), BLUE); rPr.append(colour)
            underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rPr.append(underline)
            t = OxmlElement("w:t"); t.text = piece; r.append(rPr); r.append(t); hyperlink.append(r)
            p._p.append(hyperlink)
        else:
            r = p.add_run(piece); set_font(r)
    return p


def add_pair(doc, first, second):
    path = PAIRS / f"Figures_{first:02d}_{second:02d}.jpg"
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(path), width=Inches(6.75))


def add_single(doc, number, slug):
    path = FIGURES / f"Figure_{number:02d}_{slug}.jpg"
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(path), width=Inches(6.75))


def add_bbr_pair(doc, first, second):
    path = BBR_PAIRS / f"BBR_{first:02d}_{second:02d}.jpg"
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(path), width=Inches(6.75))


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
    sec.top_margin = sec.bottom_margin = Inches(0.62)
    sec.left_margin = sec.right_margin = Inches(0.72)
    sec.header_distance = sec.footer_distance = Inches(0.32)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"; normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.12
    for sty, size, colour in [("Heading 1", 17, NAVY), ("Heading 2", 14, BLUE), ("Heading 3", 12, TEAL)]:
        s = doc.styles[sty]
        s.font.name = "Aptos Display"; s.font.size = Pt(size); s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(colour)
        s.paragraph_format.space_before = Pt(10); s.paragraph_format.space_after = Pt(5)
        s.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Imperial BBO Capstone | Component 25.2"); set_font(r, "Aptos", 8.5, MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Dr Nandakumar Theekkootu Pisharam | Successful optimisation strategies"); set_font(r, "Aptos", 8.2, MUTED)

    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10); title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Successful Optimisation Strategies"); set_font(r, "Aptos Display", 22, NAVY, True)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Stage 2: Required Capstone Component 25.2 | Section B"); set_font(r, "Aptos", 11, BLUE, True)
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("13-week black-box optimisation challenge | Eight functions | 104 prospective queries"); set_font(r, "Aptos", 9.5, MUTED)

    lines = POST.read_text(encoding="utf-8").splitlines()
    inserted = set()
    for line in lines:
        line = line.strip()
        if not line or line.startswith("# Stage") or line.startswith("## Successful"):
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=1)
            continue
        add_text(doc, line)
        if "stopped looking for one universal optimiser" in line and 1 not in inserted:
            add_pair(doc, 1, 2); inserted.add(1)
            add_pair(doc, 11, 12); inserted.add(11)
        elif "did not generalise to the rest" in line and 3 not in inserted:
            add_pair(doc, 3, 4); inserted.add(3)
            add_pair(doc, 13, 14); inserted.add(13)
            add_pair(doc, 15, 16); inserted.add(15)
            add_pair(doc, 17, 18); inserted.add(17)
            add_pair(doc, 19, 20); inserted.add(19)
        elif "stable optimum" in line and 5 not in inserted:
            add_pair(doc, 5, 6); inserted.add(5)
            add_pair(doc, 21, 22); inserted.add(21)
            add_pair(doc, 23, 24); inserted.add(23)
            add_pair(doc, 25, 26); inserted.add(25)
        elif "fixed for all 13 rounds" in line and 7 not in inserted:
            add_pair(doc, 7, 8); inserted.add(7)
            add_pair(doc, 27, 28); inserted.add(27)
            add_pair(doc, 29, 30); inserted.add(29)
        elif "overturn either explanation" in line and 9 not in inserted:
            add_pair(doc, 9, 10); inserted.add(9)
        elif "Black Box Resolution is the structured investigation" in line and 31 not in inserted:
            add_single(doc, 31, "post_bbo_resolution"); inserted.add(31)
            doc.add_page_break()
            add_bbr_pair(doc, 1, 2)
            add_bbr_pair(doc, 3, 4)

    props = doc.core_properties
    props.title = "BBO Component 25.2: Successful Optimisation Strategies"
    props.subject = "Evidence-led analysis of successful optimisation strategies and professional transfer"
    props.author = "Dr Nandakumar Theekkootu Pisharam"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
