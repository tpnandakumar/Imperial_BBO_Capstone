from pathlib import Path
import os
import re
import textwrap

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
HERE = Path(__file__).resolve().parent
ASSETS = HERE / "academic_infographics"
STORY_PAIRS = HERE / "scientific_story_pairs"
OUT = Path(os.environ.get("BBO_DISCUSSION_OUTPUT", HERE / "BBO_25_1_Academic_Retrospective_With_40_Infographics.docx"))
INCLUDE_STORIES = os.environ.get("BBO_INCLUDE_STORIES", "1") == "1"
POST = HERE / "SECTION_25_1_DISCUSSION_BOARD_POST.md"
HISTORY = ROOT / "BBO_Dashboard" / "data" / "complete_internal_evidence.csv"

NAVY = "#14213D"
BLUE = "#2F80ED"
TEAL = "#16A085"
GOLD = "#F2B134"
RED = "#D64545"
PALE = "#F4F7FB"
INK = "#263238"
MUTED = "#607080"


def save(fig, number, slug):
    path = ASSETS / f"Figure_{number:02d}_{slug}.jpg"
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    tmp = path.with_suffix(".rgb.jpg")
    with Image.open(path) as source:
        source.convert("RGB").save(tmp, format="JPEG", quality=88, optimize=True, progressive=True)
    tmp.replace(path)
    return path


def canvas(title, subtitle=""):
    fig, ax = plt.subplots(figsize=(12, 6.4))
    fig.patch.set_facecolor(PALE)
    ax.set_facecolor(PALE)
    ax.axis("off")
    ax.text(0.03, 0.93, title, transform=ax.transAxes, fontsize=25, weight="bold", color=NAVY, va="top")
    if subtitle:
        ax.text(0.03, 0.85, subtitle, transform=ax.transAxes, fontsize=12.5, color=MUTED, va="top")
    return fig, ax


def boxes(number, slug, title, subtitle, items, colours=None, columns=4):
    fig, ax = canvas(title, subtitle)
    colours = colours or [BLUE, TEAL, GOLD, RED]
    rows = int(np.ceil(len(items) / columns))
    top, bottom = 0.77, 0.10
    h = (top - bottom) / rows * 0.72
    gap_y = (top - bottom) / rows
    w = 0.90 / columns * 0.88
    for i, (label, detail) in enumerate(items):
        r, c = divmod(i, columns)
        x = 0.05 + c * (0.90 / columns)
        y = top - r * gap_y - h
        colour = colours[i % len(colours)]
        ax.add_patch(plt.Rectangle((x, y), w, h, transform=ax.transAxes, facecolor="white", edgecolor=colour, linewidth=2))
        ax.add_patch(plt.Rectangle((x, y + h - 0.045), w, 0.045, transform=ax.transAxes, facecolor=colour, edgecolor=colour))
        label_width = 16 if columns == 4 else 24
        detail_width = 21 if columns == 4 else 31
        wrapped_label = textwrap.fill(str(label), width=label_width)
        wrapped_detail = textwrap.fill(str(detail), width=detail_width)
        ax.text(x + 0.015, y + h - 0.065, wrapped_label, transform=ax.transAxes,
                fontsize=11.3 if columns == 4 else 12, weight="bold", color=NAVY, va="top", linespacing=1.05)
        label_lines = wrapped_label.count("\n") + 1
        detail_y = y + h - 0.12 - (label_lines - 1) * 0.038
        ax.text(x + 0.015, detail_y, wrapped_detail, transform=ax.transAxes,
                fontsize=9.0 if columns == 4 else 9.6, color=INK, va="top", linespacing=1.15)
    return save(fig, number, slug)


def create_infographics():
    ASSETS.mkdir(parents=True, exist_ok=True)
    history = pd.read_csv(HISTORY)
    weekly = history[history.source.str.match(r"week_\d+")].copy()
    weekly["week"] = weekly.source.str.extract(r"(\d+)").astype(int)
    figures = []

    figures.append(boxes(1, "starter_data", "The common starting point", "Course-supplied observations for eight hidden functions",
        [("F1", "2D | 10 points"), ("F2", "2D | 10 points"), ("F3", "3D | 15 points"), ("F4", "4D | 30 points"),
         ("F5", "4D | 20 points"), ("F6", "5D | 20 points"), ("F7", "6D | 30 points"), ("F8", "8D | 40 points")], columns=4))

    figures.append(boxes(2, "portal_cycle", "The weekly BBO cycle", "One new coordinate per function, followed by one returned output",
        [("1. Review", "Use every result available"), ("2. Select", "Choose eight new inputs"), ("3. Submit", "Enter values in the portal"),
         ("4. Receive", "Collect eight outputs"), ("5. Append", "Add results to the history"), ("6. Adapt", "Prepare the next week")], columns=3))

    figures.append(boxes(3, "codebase_growth", "How the codebase grew", "The analysis became more structured as evidence accumulated",
        [("Weeks 1 to 3", "Load, rank and compare"), ("Weeks 4 to 6", "Track direction and failure"), ("Weeks 7 to 9", "Validate and document"),
         ("Week 10", "Cluster recurring regions"), ("Week 11", "Test PCA structure"), ("Weeks 12 to 13", "Reward, stopping and final evidence")], columns=3))

    figures.append(boxes(4, "weeks_1_3", "Weeks 1 to 3: the functions separate", "Early results showed that one rule would not suit all eight functions",
        [("F5", "Repeated gains supported exploitation"), ("F1", "First visible positive result in Week 3"),
         ("F2, F7, F8", "Declines warned against momentum"), ("F3, F4, F6", "Weak evidence required exploration")], columns=4))

    figures.append(boxes(5, "weeks_4_6", "Weeks 4 to 6: direction matters", "The same size movement produced very different responses",
        [("Recover", "F2 and F3 improved after changing direction"), ("Extend", "F5 continued to rise"),
         ("Overreach", "F4 and F6 weakened after larger moves"), ("Learn", "Negative results narrowed the search")], columns=4))

    figures.append(boxes(6, "weeks_7_9", "Weeks 7 to 9: selective continuation", "Effort shifted towards supported regions and reproducible evidence",
        [("F5", "4278.817 to 4394.868"), ("F2 and F7", "Local regions remained promising"),
         ("F3 and F4", "Direction required reassessment"), ("Repository", "Datasheet, model card and validation added")], columns=4))

    figures.append(boxes(7, "clustering", "Week 10 clustering experiment", "Recurring local regions were compared without claiming access to the hidden surface",
        [("Inputs", "Weeks 1 to 10 coordinates and outputs"), ("Tuning", "k=2 or 3 | n_init=50 | seed=42"),
         ("Interpretation", "Productive, weak and plateau regions"), ("Decision", "Recovery, refinement or confirmation for Week 11")], columns=4))

    figures.append(boxes(8, "pca", "Week 11 PCA analysis", "PCA described the path taken through higher-dimensional spaces",
        [("Variance", "How concentrated was the search path?"), ("Correlation", "Which coordinates moved together?"),
         ("Redundancy", "Did several coordinates carry similar movement?"), ("Boundary", "PCA supported decisions only when outputs agreed")], columns=4))

    figures.append(boxes(9, "week12_state", "Week 12 evidence before the final query", "Different functions reached different decision states",
        [("New best", "F2, F3 and F5"), ("Winner repeated", "F1, F4, F7 and F8"),
         ("Uncertainty", "F6 did not reproduce its earlier reward"), ("Final question", "Move, retain, recover or repeat?")], columns=4))

    figures.append(boxes(10, "rl_mapping", "Week 13 reinforcement learning analysis", "The 12-week history was mapped into a state-action-reward decision",
        [("State", "Accumulated evidence for each function"), ("Action", "The next submitted coordinate"),
         ("Reward", "The returned black-box output"), ("Policy", "Explore, refine, retain or repeat")], columns=4))

    figures.append(boxes(11, "final_actions", "Week 13 action allocation", "The final round used four function-specific actions",
        [("Retain winners", "F1, F4, F7 and F8"), ("Local refinement", "F2 and F3"),
         ("Boundary refinement", "F5"), ("Repeat for uncertainty", "F6")], columns=4))

    final = pd.DataFrame({"Function":[f"F{i}" for i in range(1,9)],
        "Output":[0.0255592853,0.7335252043,-0.0568506160,-4.3598749266,4440.9572166,-0.6071562249,1.3809299934,9.58024],
        "Week":["3, 11, 12, 13","12","13","1, 12, 13","13","13","5, 12, 13","1, 11, 12, 13"]})
    fig, ax = canvas("Strongest verified results", "Best observed value and the week or weeks in which it occurred")
    ax.axis("off")
    y = 0.76
    for i,row in final.iterrows():
        yy = y - i*0.082
        ax.add_patch(plt.Rectangle((0.05,yy),0.90,0.06,transform=ax.transAxes,facecolor="white",edgecolor="#DCE4EE"))
        ax.text(0.07,yy+0.03,row.Function,transform=ax.transAxes,va="center",weight="bold",color=BLUE,fontsize=12)
        ax.text(0.20,yy+0.03,f"{row.Output:.10g}",transform=ax.transAxes,va="center",color=INK,fontsize=11)
        ax.text(0.57,yy+0.03,f"Best week(s): {row.Week}",transform=ax.transAxes,va="center",color=MUTED,fontsize=11)
    figures.append(save(fig,12,"final_results"))

    f5 = weekly[weekly.function==5].sort_values("week")
    fig, ax = plt.subplots(figsize=(12,6.4)); fig.patch.set_facecolor(PALE); ax.set_facecolor("white")
    ax.plot(f5.week, f5.output, marker="o", linewidth=3, color=TEAL)
    ax.set_title("F5: sustained improvement towards the boundary", loc="left", fontsize=23, weight="bold", color=NAVY)
    ax.set_xlabel("Week"); ax.set_ylabel("Returned output"); ax.grid(alpha=.2); ax.spines[["top","right"]].set_visible(False)
    ax.annotate("1415.876",(1,f5.iloc[0].output),xytext=(1.3,1800),arrowprops={"arrowstyle":"->"})
    ax.annotate("4440.957",(13,f5.iloc[-1].output),xytext=(10.4,4100),arrowprops={"arrowstyle":"->"})
    figures.append(save(fig,13,"f5_trajectory"))

    figures.append(boxes(14, "f2_overshoot", "F2: a small move can still overshoot", "Week 12 remained the winner after a 0.005 change in the first coordinate",
        [("Week 12", "0.690000, 0.950000\nReward 0.733525"), ("Week 13", "0.685000, 0.950000\nReward 0.641343"),
         ("Movement", "L1 distance 0.005000"), ("Change", "Reward fell by 0.092182")], columns=4))

    figures.append(boxes(15, "f3_refinement", "F3: controlled local refinement succeeded", "A small three-coordinate movement produced a new overall best",
        [("Week 12", "0.850, 0.150, 0.850"), ("Week 13", "0.855, 0.145, 0.855"),
         ("Movement", "L1 distance 0.015000"), ("Gain", "+0.003000659")], columns=4))

    figures.append(boxes(16, "f6_variability", "F6: identical coordinate, different rewards", "Repeated evaluation exposed unresolved response variability",
        [("Coordinate", "0.700, 0.200, 0.700, 0.700, 0.200"), ("Week 3", "-0.648848"),
         ("Week 12", "-0.707832"), ("Week 13", "-0.607156 | new best")], columns=4))

    figures.append(boxes(17, "recovery_retention", "Recovery and retention protected earlier winners", "Not every successful final decision required further movement",
        [("F1", "Winner confirmed in Weeks 3, 11, 12, 13"), ("F4", "Week 1 winner recovered in Weeks 12 and 13"),
         ("F7", "Week 5 winner recovered in Weeks 12 and 13"), ("F8", "Week 1 winner confirmed three more times")], columns=4))

    figures.append(boxes(18, "four_actions", "The four-action search framework", "A practical vocabulary that emerged from the weekly evidence",
        [("Explore", "Test a different region"), ("Refine", "Move carefully near a promising point"),
         ("Recover", "Return towards an earlier winner"), ("Retain", "Stop moving when the result is confirmed")], columns=4))

    figures.append(boxes(19, "stopping", "Where optimisation would continue", "Stopping was decided separately for each function",
        [("Continue", "F2: recover the Week 12 region"), ("Continue", "F3: final direction still improved"),
         ("Continue", "F5: boundary trend remained productive"), ("Continue", "F6: uncertainty remained unresolved"),
         ("Stop routine search", "F1, F4, F7 and F8"), ("Reason", "Winners had been repeatedly confirmed")], columns=3))

    figures.append(boxes(20, "clinical_application", "Application to clinical neurology", "The capstone lessons transfer to cautious service improvement",
        [("Preserve provenance", "Know where every decision came from"), ("Test sequentially", "Change one pathway step at a time"),
         ("Protect high-risk cases", "Do not trade safety for average gain"), ("Check variation", "Separate improvement from noise"),
         ("Use stopping rules", "Stop when evidence supports stopping"), ("State limits", "Be clear about what data cannot prove")], columns=3))
    return figures


def set_font(run, name="Aptos", size=10.5, colour=INK, bold=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(colour.replace("#", ""))
    run.bold = bold


def add_rich_text(paragraph, text):
    pieces = re.split(r"(`[^`]+`|\[[^\]]+\]\([^)]+\))", text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("`"):
            run = paragraph.add_run(piece[1:-1]); set_font(run, "Consolas", 9.5, NAVY)
        elif piece.startswith("["):
            label = piece[1:piece.index("]")]
            run = paragraph.add_run(label); set_font(run, "Aptos", 10.5, BLUE)
            run.underline = True
        else:
            run = paragraph.add_run(piece); set_font(run)


def build_doc(figures):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
    sec.top_margin = sec.bottom_margin = Inches(0.55)
    sec.left_margin = sec.right_margin = Inches(0.72)
    sec.header_distance = sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.16
    for sty, size, colour in [("Heading 1",17,NAVY),("Heading 2",14,BLUE),("Heading 3",12,TEAL)]:
        s=doc.styles[sty]; s.font.name="Aptos Display"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(colour.replace("#",""))
        s.paragraph_format.space_before=Pt(10); s.paragraph_format.space_after=Pt(5); s.paragraph_format.keep_with_next=True

    title = doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=title.add_run("Retrospective on the BBO Capstone Project"); set_font(r,"Aptos Display",22,NAVY,True)
    sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=sub.add_run("Stage 2: Required Capstone Component 25.1 | Section B"); set_font(r,"Aptos",11,BLUE,True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Dr Nandakumar Theekkootu Pisharam | 13-week BBO challenge"); set_font(r,"Aptos",9.5,MUTED)

    text = POST.read_text(encoding="utf-8").splitlines()
    content=[]; in_table=False
    for line in text:
        if line.startswith("# Stage") or line.startswith("## Retrospective"):
            continue
        if line.startswith("|"):
            in_table=True; continue
        if in_table and not line.startswith("|"):
            in_table=False
        if line.startswith("![") or line.startswith("*Figure"):
            continue
        content.append(line)

    placement={
        (1,1):[1], (1,3):[2],
        (2,1):[3,4], (2,2):[13,14], (2,3):[5,6],
        (2,4):[15,16], (2,5):[17,18], (2,6):[11,12], (2,7):[19],
        (3,1):[20], (3,2):[9], (3,3):[7], (3,4):[10],
        (4,3):[8],
    }
    story_placement={
        (1,1):1, (1,3):2,
        (2,1):3, (2,2):4, (2,3):5, (2,4):6,
        (3,1):7, (3,3):8,
        (4,1):9,
        (5,1):10,
    }

    def add_figures(numbers):
        for n in numbers:
            pic=doc.add_paragraph(); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER; pic.paragraph_format.keep_with_next=True
            pic.paragraph_format.space_before=Pt(3); pic.paragraph_format.space_after=Pt(1)
            pic.add_run().add_picture(str(figures[n-1]), width=Inches(6.65))
            cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after=Pt(7)
            rr=cap.add_run(f"Figure {n}"); set_font(rr,"Aptos",8.5,MUTED,True)

    def add_story_pair(pair_number):
        first = pair_number * 2 - 1
        path = STORY_PAIRS / f"Stories_{first:02d}_{first+1:02d}.jpg"
        pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER; pic.paragraph_format.keep_with_next = True
        pic.paragraph_format.space_before = Pt(4); pic.paragraph_format.space_after = Pt(1)
        pic.add_run().add_picture(str(path), width=Inches(6.78))
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_after = Pt(8)
        rr = cap.add_run(f"Scientific story infographics S{first:02d} and S{first+1:02d}")
        set_font(rr,"Aptos",8.5,MUTED,True)
    current=0
    paragraph_number=0
    for line in content:
        if not line.strip():
            continue
        if line.startswith("### "):
            current=int(line.split()[1].rstrip("."))
            paragraph_number=0
            doc.add_heading(line[4:], level=1)
        else:
            paragraph_number += 1
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.widow_control=True
            add_rich_text(p,line)
            add_figures(placement.get((current,paragraph_number),[]))
            pair_number = story_placement.get((current,paragraph_number)) if INCLUDE_STORIES else None
            if pair_number:
                add_story_pair(pair_number)

    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=footer.add_run("Imperial BBO Capstone | Component 25.1"); set_font(rr,"Aptos",8,MUTED)
    doc.core_properties.title="BBO Capstone Component 25.1 Discussion"
    doc.core_properties.author="Dr Nandakumar Theekkootu Pisharam"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    from generate_academic_infographics import generate
    figs=generate()
    output=build_doc(figs)
    print(output)
